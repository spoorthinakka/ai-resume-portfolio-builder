"""
AI Resume & Portfolio Builder (Streamlit)

Purpose
-------
Generate a clean, professional resume from user inputs using a Hugging Face model,
and export it as PDF/DOCX/TXT. Optionally builds a simple portfolio website (HTML zip).

Key Features
------------
- Streamlit UI for contact, education, skills, projects, internships/experience
- LLM drafting via Hugging Face Inference API
- Smart post-processing (headings, bullets, section order)
- Downloads: PDF / DOCX / TXT; Portfolio .zip (HTML+CSS)
- Environment-first config (no secrets in code)

Environment Variables
---------------------
HF_TOKEN : Hugging Face Inference API token (required)
HF_MODEL : Optional model id (defaults to meta-llama/Llama-3.1-8B-Instruct, or any compatible chat-instruct model you prefer)

Examples
--------
HF_MODEL could be:
- meta-llama/Llama-3.1-8B-Instruct  (default)
- mistralai/Mistral-Nemo-Instruct-2407
- Qwen/Qwen2.5-7B-Instruct
(Use an instruct/chat model for best results.)

Security Note
-------------
Never commit tokens. In PowerShell:  $env:HF_TOKEN="hf_xxx"
"""


import os, io, json, re, zipfile
from io import BytesIO
from datetime import datetime
import streamlit as st
from huggingface_hub import InferenceClient
from dotenv import load_dotenv

# Load environment variables (API Key that you've created)
load_dotenv()

# ====================== SETUP & CONFIG ======================

# Configure Streamlit page and title
st.set_page_config(page_title="AI Resume & Portfolio Builder", page_icon="🧠", layout="centered")
st.title("🧠 AI Resume & Portfolio Builder")

# Read credentials/config from environment (keep secrets out of code)
HF_TOKEN = os.getenv("HF_TOKEN")   # e.g., PowerShell: $env:HF_TOKEN="hf_xxx"
HF_MODEL = "meta-llama/Llama-3.1-8B-Instruct"  # allow override via env

# Initialize HF Inference client only if a token is present
# (Prevents crashes on startup; we still check again before calling the API.)
client = InferenceClient(HF_MODEL, token=HF_TOKEN) 

# ====================== SIDEBAR ======================
#Sidebar handles resume style options, quick visual previews, and portfolio theme choice.

with st.sidebar:
    # ----- Template selection -----
    st.header("🎨 Template")
    template = st.selectbox("Style", ["Modern", "Classic", "Professional"], index=0)

    # ----- Generation options -----
    st.header("⚙️ Options")
    ai_overview = st.checkbox("Let AI write Professional Overview", value=True)

# Define preview placeholders so sidebar updates live even before form submission
title_preview = (st.session_state.get("pro_title") or "Software Engineer")
name_preview = (st.session_state.get("name") or "NAME SURNAME").upper()
loc_preview = (st.session_state.get("location") or "City, Country")

# ----- TEMPLATE PREVIEWS (visual difference among Modern / Classic / Professional) -----
with st.sidebar:
    st.header("🪟 Template Preview (styled)")

    if template == "Modern":
        st.markdown(f"""
        <style>
          /* Make colors readable on dark sidebars */
          .prev-modern {{ color:#E5E7EB; }}
          .prev-modern a {{ color:#E5E7EB; }}
          .prev-modern h1 {{ font-family: Inter, system-ui, sans-serif; font-weight:800; font-size:16px; color:#E5E7EB; margin:0; text-transform:uppercase; }}
          .prev-modern .sub {{ color:#CBD5E1; font-size:12px; }}
          .prev-modern .chip {{ display:inline-block; background:#F8FAFC; color:#111827; border:1px solid #E2E8F0; border-radius:9999px; padding:2px 8px; font-size:11px; font-weight:600; margin-top:6px; }}
          .prev-modern hr {{ border:0; border-top:2px solid rgba(229,231,235,.45); margin:6px 0; }}
          .prev-modern .card {{ background:#0B1220; border-radius:14px; box-shadow:0 6px 24px rgba(0,0,0,.35); padding:10px; color:#E5E7EB; border:1px solid rgba(229,231,235,.18); }}
        </style>
        <div class="prev-modern" style="font-family:Inter,system-ui,sans-serif; line-height:1.35;">
          <div style="text-align:center;">
            <h1>{name_preview}</h1>
            <div class="sub">{loc_preview}</div>
            <div class="sub">Email | Phone · LinkedIn | GitHub</div>
            <div class="chip">{title_preview}</div>
          </div>
          <div style="margin-top:10px; font-weight:700; font-size:12px; color:#E5E7EB; letter-spacing:.05em;">PROFESSIONAL OVERVIEW</div>
          <hr>
          <div class="card" style="font-size:12px;">Clean, modern layout with neutral accents and soft cards.</div>
        </div>
        """, unsafe_allow_html=True)

    elif template == "Classic":
        st.markdown(f"""
        <div style="font-family: 'Times New Roman', Times, serif; line-height:1.35; color:#E5E7EB;">
          <div style="text-align:center;">
            <div style="font-weight:800; font-size:16px; color:#FFFFFF;">{name_preview} — {title_preview}</div>
            <div style="color:#CBD5E1; font-size:12px;">{loc_preview}</div>
            <div style="font-size:12px; color:#E5E7EB;">Email | Phone</div>
            <div style="font-size:12px; color:#E5E7EB;">LinkedIn | GitHub</div>
          </div>
          <div style="margin-top:10px; font-weight:800; font-size:12px; color:#FFFFFF;">PROFESSIONAL OVERVIEW</div>
          <hr style="border: 1px solid rgba(229,231,235,.6); margin:6px 0;">
          <div style="font-size:12px; color:#E5E7EB;">Traditional serif with thin black dividers. Formal & print-friendly.</div>
        </div>
        """, unsafe_allow_html=True)

    else:  # Professional – professional, minimal, monochrome
        st.markdown(f"""
        <style>
          .prev-professional {{ color:#E5E7EB; }}
          .prev-professional h1 {{ font-family: 'Playfair Display', serif; font-weight:700; font-size:16px; color:#FFFFFF; margin:0; letter-spacing:.2px; }}
          .prev-professional .sub {{ color:#CBD5E1; font-size:12px; }}
          .prev-professional hr {{ border:0; border-top:1px solid rgba(229,231,235,.45); margin:6px 0; }}
          .prev-professional .blurb {{ font-size:12px; color:#E5E7EB; }}
        </style>
        <link href="https://fonts.googleapis.com/css2?family=Playfair+Display:wght@600;700&display=swap" rel="stylesheet">
        <div class="prev-professional" style="font-family: Inter, system-ui, -apple-system, Segoe UI, Roboto, sans-serif; line-height:1.5;">
          <div style="text-align:center;">
            <h1>{name_preview}</h1>
            <div class="sub">{loc_preview}</div>
            <div class="sub">Email | Phone · LinkedIn | GitHub</div>
          </div>
          <div style="margin-top:10px; font-weight:700; font-size:11px; letter-spacing:.08em; text-transform:uppercase; color:#E5E7EB;">Professional Overview</div>
          <hr>
          <div class="blurb">Professional typography with spacious layout and subtle contrasts.</div>
        </div>
        """, unsafe_allow_html=True)

    st.header("🌐 Portfolio Theme")
    portfolio_theme = st.selectbox("Website Theme", ["Modern", "Professional"], index=0)

# ====================== FORM: Collect user inputs for resume generation ======================
# Captures all information needed by the model and exporters.
# Sections: Contact → Education → Skills → Experience → Projects → Publications → Extras

with st.form("resume_form"):
    st.subheader("Contact Information")
    name = st.text_input("Full Name", key="name")
    pro_title = st.text_input("Professional Title (e.g., Developer, Data Analyst, AI/ML Engineer, UX Designer)", key="pro_title")
    location = st.text_input("Location (City, State, Country)", key="location")
    email = st.text_input("Email", key="email")
    phone = st.text_input("Phone Number", key="phone")
    linkedin = st.text_input("LinkedIn Profile", key="linkedin")
    github = st.text_input("GitHub Profile", key="github")

    st.subheader("Professional Overview")
    summary = st.text_area("Provide a short professional summary (2–4 lines) highlighting your background, skills, and career goals. You may skip this section if you prefer the AI to generate it for you.", key="summary")

    # ----- Education Section -----
    st.subheader("Education")
    # Capture one education entry
    def edu_block(num: int):
        # Collect institute, degree, cgpa, and dates; return as dict.
        st.markdown(f"**Education #{num}**")
        inst   = st.text_input("Institute / University", key=f"edu_inst_{num}")
        degree = st.text_input("Degree (e.g., B.Tech in CSE, B.A, B.Com)", key=f"edu_deg_{num}")
        cgpa   = st.text_input("CGPA/Percentage", key=f"edu_cgpa_{num}")
        c1, c2 = st.columns(2)
        start  = c1.text_input("Start (e.g., Aug 2022)", key=f"edu_start_{num}")
        end    = c2.text_input("End (e.g., May 2026)", key=f"edu_end_{num}")
        return {"inst": inst, "degree": degree, "cgpa": cgpa, "start": start, "end": end}
    edu1 = edu_block(1); st.divider()
    edu2 = edu_block(2)

    # ----- Skills Section -----
    st.subheader("Skills")
    skills_lang  = st.text_area("Programming Languages (e.g., Python, C, Java)", key="skills_lang")
    skills_fw    = st.text_area("Frameworks / Tools (e.g., Django, React, Git)", key="skills_fw")
    skills_db    = st.text_area("Databases (e.g., MySQL, MongoDB)", key="skills_db")
    skills_cloud = st.text_area("Cloud / DevOps (e.g., AWS, Docker)", key="skills_cloud")
    skills_soft  = st.text_area("Soft Skills (e.g., Communication, Leadership)", key="skills_soft")

    # ----- Experience / Internships Section -----
    st.subheader("Experience / Internships")
    # Capture one experience entry
    def exp_block(num: int):
        # Collect organization, role, period, and key deliverables; return as dict.
        st.markdown(f"**Experience #{num}**")
        comp   = st.text_input("Organization / Company", key=f"exp_comp_{num}")
        role   = st.text_input("Role", key=f"exp_role_{num}")
        c1, c2 = st.columns(2)
        start  = c1.text_input("Start (e.g., May 2025)", key=f"exp_start_{num}")
        end    = c2.text_input("End (e.g., Nov 2025)", key=f"exp_end_{num}")
        desc   = st.text_area("Key Deliverables", key=f"exp_desc_{num}")
        return {"company": comp, "role": role, "start": start, "end": end, "desc": desc}
    exp1 = exp_block(1); st.divider()
    exp2 = exp_block(2); st.divider()
    exp3 = exp_block(3)

    # ----- Projects Section -----
    st.subheader("Projects")
    # Capture one project entry
    def proj_block(num: int):
        # Collect title, authors, venue, year, summary, link; return as dict.
        st.markdown(f"**Project #{num}**")
        title   = st.text_input("Title", key=f"proj_title_{num}")
        problem = st.text_area("Problem", key=f"proj_prob_{num}")
        approach= st.text_area("Approach", key=f"proj_app_{num}")
        tech    = st.text_area("Tech / Tools Used", key=f"proj_tech_{num}")
        impact  = st.text_area("Impact / Result", key=f"proj_imp_{num}")
        link    = st.text_input("Project Link (optional, e.g., GitHub repo/demo)", key=f"proj_link_{num}")
        return {"title": title, "problem": problem, "approach": approach, "tech": tech, "impact": impact, "link": link}
    proj1 = proj_block(1); st.divider()
    proj2 = proj_block(2); st.divider()
    proj3 = proj_block(3)

    # ----- Publications Section (optional) -----
    st.subheader("Publications (optional)")
    # Capture one publication entry
    def pub_block(num: int):
        # Collect title, authors, venue, year, summary, link; return as dict.
        st.markdown(f"**Publication #{num}**")
        title  = st.text_input("Title", key=f"pub_title_{num}")
        authors= st.text_input("Authors", key=f"pub_authors_{num}")
        venue  = st.text_input("Venue / Conference / Journal", key=f"pub_venue_{num}")
        year   = st.text_input("Year", key=f"pub_year_{num}")
        summary= st.text_area("Summary (1–2 lines)", key=f"pub_sum_{num}")
        link   = st.text_input("Link (optional)", key=f"pub_link_{num}")
        return {"title": title, "authors": authors, "venue": venue, "year": year, "summary": summary, "link": link}
    pub1 = pub_block(1); st.divider()
    pub2 = pub_block(2)

    # ----- Certifications / Hands-on Section -----
    st.subheader("Certifications / Hands-on")
    # Collect up to 5 certifications.
    certs = [st.text_input(f"Certification #{i}", key=f"cert_{i}") for i in range(1, 6)]

    # ----- Achievements Section -----
    st.subheader("Achievements (optional)")
    # Collect up to 5 achievements.
    achieves = [st.text_input(f"Achievement #{i}", key=f"ach_{i}") for i in range(1, 6)]

    # ----- Participations Section -----
    st.subheader("Participations (optional)")
    # Collect up to 5 participation entries.
    parts = [st.text_input(f"Participation #{i}", key=f"part_{i}") for i in range(1, 6)]

    # ----- Positions of Responsibility / Co-curricular Involvement -----
    st.subheader("Positions of Responsibility / Co-curricular Involvement (optional)")
    # Capture one position entry
    def por_block(num: int):
        # Collect role, organization, duration, responsibilities; return as dict.
        st.markdown(f"**Position #{num}**")
        role = st.text_input("Role / Title", key=f"por_role_{num}")
        org  = st.text_input("Organization", key=f"por_org_{num}")
        when = st.text_input("Period (e.g., 2023–2024)", key=f"por_when_{num}")
        det  = st.text_area("Key Responsibilities", key=f"por_det_{num}")
        return {"role": role, "org": org, "when": when, "det": det}
    por1 = por_block(1); st.divider()
    por2 = por_block(2); st.divider()
    por3 = por_block(3)

    # ----- Target Role & Job Description Section -----
    st.subheader("Target Role (optional)")
    target_role = st.text_input("Desired Position or Target Role", key="target_role")

    st.subheader("Job Description (optional)")
    job_desc = st.text_area("Paste JD to tailor your Resume / CV ", key="job_desc")

    # Finalize form and trigger generation
    submitted = st.form_submit_button("Generate Resume")

# ====================== TEXT HELPERS ======================
# These small utilities clean, normalize, and post-process text before export or model prompts.
# They make sure all generated resumes follow consistent formatting and bullet rules.

HEADINGS = {
    "PROFESSIONAL OVERVIEW","EDUCATION","SKILLS","EXPERIENCE / INTERNSHIPS",
    "PROJECTS","PUBLICATIONS","CERTIFICATIONS / HANDS-ON","ACHIEVEMENTS","PARTICIPATIONS",
    "POSITIONS OF RESPONSIBILITY / CO-CURRICULAR INVOLVEMENT","TARGET ROLE"
}

# Normalize a heading: trim spaces, remove colons, and uppercase.
def normalize_heading(line: str) -> str:
    s = line.strip()
    if s.endswith(":"): s = s[:-1]
    return s.upper()

# Sanitize text: replace smart quotes/dashes, collapse spaces, remove non-ASCII chars.
def sanitize_plain_text(text: str) -> str:
    replacements = {"•":"-","–":"-","—":"-","’":"'","‘":"'","“":'"',"”":'"',"\u00A0":" "}
    for bad, good in replacements.items():
        text = text.replace(bad, good)
    lines = [" ".join(ln.split()) for ln in text.splitlines()]
    text = "\n".join(lines)
    return "".join(ch if ord(ch) < 128 else " " for ch in text)

# Detect if a line looks like a project/publication title (no colon or dash).
def is_project_or_pub_title_line(line: str) -> bool:
    return bool(line.strip()) and (":" not in line) and (not line.strip().startswith("-"))

# Remove unwanted heading/menu lines before main resume content.
def strip_heading_menu(text: str) -> str:
    lines = [l for l in text.splitlines()]
    i = 0; saw_non_heading = False
    for j, l in enumerate(lines):
        t = l.strip()
        if not t: continue
        if normalize_heading(t) in HEADINGS:
            i = j + 1; continue
        else:
            saw_non_heading = True; break
    return "\n".join(lines[i:]).strip() if saw_non_heading else ""

# Ensure first section starts with "PROFESSIONAL OVERVIEW" (adds it if missing).
def ensure_first_section_heading(text: str, heading="PROFESSIONAL OVERVIEW") -> str:
    lines = [l.rstrip() for l in text.splitlines()]
    i = 0
    while i < len(lines) and not lines[i].strip():
        i += 1
    if i == len(lines):
        return f"{heading}\n"
    first = lines[i].strip().rstrip(":").upper()
    if first != heading:
        lines.insert(i, heading)
        if i + 1 < len(lines) and not lines[i + 1].strip():
            del lines[i + 1]
    else:
        if i + 1 < len(lines) and not lines[i + 1].strip():
            del lines[i + 1]
    return "\n".join(lines)

# Enforce bullet markers in list-style sections (skills, certs, achievements, etc.).    
def enforce_bullets_in_sections(text: str) -> str:
    """Force bullets in these sections even if the LLM outputs plain lines."""
    BULLET_SECTIONS = {
        "SKILLS",
        "PUBLICATIONS",
        "CERTIFICATIONS / HANDS-ON",
        "ACHIEVEMENTS",
        "PARTICIPATIONS",
    }
    lines = text.splitlines()
    out, in_block = [], False
    for raw in lines:
        ln = raw.rstrip()
        up = normalize_heading(ln)

        if up in BULLET_SECTIONS:
            in_block = True
            out.append(up)
            continue

        if in_block and up in HEADINGS:
            in_block = False
            out.append(up)
            continue

        if in_block:
            s = ln.strip()
            if s and not s.startswith("- "):
                ln = "- " + s
        out.append(ln)

    return "\n".join(out)

# ====================== FORMATTERS ======================
# Format small data dicts into compact resume-ready lines/blocks.

# Join only non-empty strings with a separator (used for dates/contacts).
def join_nonempty(vals, sep=" | "):
    return sep.join([v for v in vals if v and v.strip()])

# Format one education entry → "Institute — Degree | CGPA X | Start – End".
def fmt_edu(e):
    if not any(e.values()): return ""
    bits = []
    if e["inst"] and e["degree"]: bits.append(f"{e['inst']} — {e['degree']}")
    elif e["inst"]: bits.append(e["inst"])
    elif e["degree"]: bits.append(e["degree"])
    if e["cgpa"]:  bits.append(f"CGPA {e['cgpa']}")
    date_line = join_nonempty([e["start"], e["end"]], " – ")
    if date_line: bits.append(date_line)
    return " | ".join(bits)

# Format one experience entry → "Company — Role | Period" + bullet for deliverables.
def fmt_exp(e):
    if not any(e.values()): return ""
    when = join_nonempty([e["start"], e["end"]], " – ")
    header = join_nonempty([e["company"], e["role"]], " — ")
    line1 = f"{header} | {when}" if when else header
    out = [line1]
    if e["desc"]: out.append(f"- {e['desc']}")
    return "\n".join(out)

# Format one project → title + bullets for Problem/Approach/Tech/Impact/Link.
def fmt_proj(p):
    if not any(p.values()): return ""
    lines = []
    if p["title"]:   lines.append(p["title"])
    if p["problem"]: lines.append(f"- Problem: {p['problem']}")
    if p["approach"]:lines.append(f"- Approach: {p['approach']}")
    if p["tech"]:    lines.append(f"- Tech/Tools: {p['tech']}")
    if p["impact"]:  lines.append(f"- Impact: {p['impact']}")
    if p.get("link"): lines.append(f"- Link: {p['link']}")
    return "\n".join(lines)

# Format one publication → title + bullets for Authors/Venue-Year/Summary/Link.
def fmt_pub(p):
    if not any(p.values()): return ""
    parts = []
    if p["title"]: parts.append(p["title"])
    if p["authors"]: parts.append(f"- Authors: {p['authors']}")
    if p["venue"] or p["year"]:
        v = join_nonempty([p["venue"], p["year"]], ", ")
        parts.append(f"- Venue/Year: {v}")
    if p["summary"]: parts.append(f"- Summary: {p['summary']}")
    if p["link"]: parts.append(f"- Link: {p['link']}")
    return "\n".join(parts)

# Format one position of responsibility → header + bullet for details.
def fmt_por(p):
    if not any(p.values()): return ""
    head = join_nonempty([p["role"], p["org"], p["when"]], " | ")
    lines = [head]
    if p["det"]: lines.append(f"- {p['det']}")
    return "\n".join(lines)

# Build the SKILLS block from the various skills text areas.
def skills_block():
    items = []
    if skills_lang.strip():  items.append(f"- Programming Languages: {skills_lang.strip()}")
    if skills_fw.strip():    items.append(f"- Frameworks/Tools: {skills_fw.strip()}")
    if skills_db.strip():    items.append(f"- Databases: {skills_db.strip()}")
    if skills_cloud.strip(): items.append(f"- Cloud/DevOps: {skills_cloud.strip()}")
    if skills_soft.strip():  items.append(f"- Soft Skills: {skills_soft.strip()}")
    return "\n".join(items)

# Build the 1–4 header lines (name/title, location, contacts, links).
def build_header_lines():
    h1 = f"{(name or '').upper()} — {pro_title}"
    h2 = location.strip() if location.strip() else ""
    h3 = join_nonempty([f"Email: {email}", f"Phone: {phone}"], " | ")
    h4 = join_nonempty([f"LinkedIn: {linkedin}", f"GitHub: {github}"], " | ")
    return h1, h2, h3, h4

# ====================== PROMPT DATA ======================
# Build model-ready profile text and a strict, deterministic prompt.

# Assemble all form inputs into a clean, sectioned profile text for the model.
def build_profile_text():
    edu_lines = "\n".join([x for x in [fmt_edu(edu1), fmt_edu(edu2)] if x])
    skills_txt = skills_block()
    exp_lines = "\n\n".join([x for x in [fmt_exp(exp1), fmt_exp(exp2), fmt_exp(exp3)] if x])
    proj_lines= "\n\n".join([x for x in [fmt_proj(proj1), fmt_proj(proj2), fmt_proj(proj3)] if x])

    pub_lines = "\n\n".join([x for x in [fmt_pub(pub1), fmt_pub(pub2)] if x])

    cert_lines = "\n".join([f"- {c.strip()}" for c in certs if c.strip()])
    ach_lines  = "\n".join([f"- {a.strip()}" for a in achieves if a.strip()])
    part_lines = "\n".join([f"- {p.strip()}" for p in parts if p.strip()])

    por_lines  = "\n\n".join([x for x in [fmt_por(por1), fmt_por(por2), fmt_por(por3)] if x])

    overview_block = "" if ai_overview else summary.strip()

    return f"""
PROFESSIONAL OVERVIEW:
{overview_block}

EDUCATION:
{edu_lines}

SKILLS:
{skills_txt}

EXPERIENCE / INTERNSHIPS:
{exp_lines}

PROJECTS:
{proj_lines}

PUBLICATIONS:
{pub_lines}

CERTIFICATIONS / HANDS-ON:
{cert_lines}

ACHIEVEMENTS:
{ach_lines}

PARTICIPATIONS:
{part_lines}

POSITIONS OF RESPONSIBILITY / CO-CURRICULAR INVOLVEMENT:
{por_lines}

TARGET ROLE:
{target_role}
"""

# Create the final LLM prompt (order, rules, and optional JD tailoring).
def build_prompt(profile_text: str) -> str:
    jd_block = f"\nJob Description to tailor to:\n{job_desc}\n" if job_desc.strip() else ""
    ai_overview_rule = (
        "\n- WRITE the 'PROFESSIONAL OVERVIEW' as a comprehensive 5–7 sentence CV-style summary using the profile below. "
        "Make it tailored and impactful (projects, tools, outcomes)."
        if ai_overview else
        "\n- Use the given 'PROFESSIONAL OVERVIEW' content as-is."
    )
    return f"""
Output CLEAN PLAIN TEXT (no markdown). Do not add any preface text.
Begin directly with the heading 'PROFESSIONAL OVERVIEW' in uppercase.

Order and headings (uppercase). Do NOT include a 'CV SUMMARY' section:
PROFESSIONAL OVERVIEW
EDUCATION
SKILLS
EXPERIENCE / INTERNSHIPS
PROJECTS
PUBLICATIONS
CERTIFICATIONS / HANDS-ON
ACHIEVEMENTS
PARTICIPATIONS
POSITIONS OF RESPONSIBILITY / CO-CURRICULAR INVOLVEMENT

Rules:
- Do not invent info. Omit empty sections.
- Use bullets only inside sections where the source text already has list-like lines.
- In PROJECTS/PUBLICATIONS, keep structured sub-lines (Problem/Approach/Tech/Impact, Authors/Venue/Year/Summary/Link).
- Include link lines if present (e.g., '- Link: ...').{ai_overview_rule}
{jd_block}
Here is the profile:
{profile_text}
"""

# ====================== GENERATION ======================
# Validate inputs, build prompt, call HF API, normalize output, and stash in session.

if submitted:
    # Guard: basic sanity checks before model call
    if not name or not email:
        st.warning("Please enter at least your Name and Email.")
    elif not HF_TOKEN:
        st.error("Missing HF_TOKEN environment variable.")
    else:
        with st.spinner("Generating your resume..."):
            try:
                # 1) Build profile text from form inputs
                profile = build_profile_text()
                
                # 2) Create deterministic prompt (order + rules)
                prompt = build_prompt(profile)
                
                # 3) Call Hugging Face chat-instruct model
                resp = client.chat_completion(
                    messages=[
                        {"role": "system", "content": "Be concise and format professional resumes in plain text."},
                        {"role": "user", "content": prompt}
                    ],
                    max_tokens=1600,
                    temperature=0.6,
                )
                
                # 4) Normalize model output (strip menus, enforce first heading, bullets)
                body = resp.choices[0].message["content"].strip()
                body = strip_heading_menu(body)
                body = ensure_first_section_heading(body, "PROFESSIONAL OVERVIEW")
                body = enforce_bullets_in_sections(body)  

                # 5) Prepend contact header block
                h1, h2, h3, h4 = build_header_lines()
                header_lines = [x for x in [h1, h2, h3, h4] if x]
                final_text = f"{'\n'.join(header_lines)}\n\n{body}"

                # 6) Save to session for editor/scoring/exports
                st.session_state["final_text"] = final_text
                st.session_state["editor"] = final_text
                st.success("Generated successfully. You can edit below and save before downloading.")
            
            except Exception as e:
                # Fallback: show root cause without crashing the app
                st.error(str(e))

# ====================== SUGGESTIONS (sidebar) ======================
# Lightweight heuristics to nudge users to complete key fields before export/portfolio.

def build_suggestions():
    # Build a short list of actionable tips based on missing fields.
    tips = []
    if not st.session_state.get("linkedin","").strip():
        tips.append("Add your LinkedIn URL for credibility.")
    if not st.session_state.get("github","").strip():
        tips.append("Add your GitHub URL if you have projects/code samples.")
    if not st.session_state.get("summary","").strip() and not ai_overview:
        tips.append("Write a 2–4 line Professional Overview or enable AI.")
    if not st.session_state.get("skills_lang","").strip(): tips.append("List 3–6 programming languages.")
    if not st.session_state.get("skills_fw","").strip():   tips.append("Add key frameworks/tools (Django/React/Git etc.)")
    if not st.session_state.get("skills_db","").strip():   tips.append("Mention at least one database (MySQL/MongoDB).")
    if not st.session_state.get("edu_inst_1","").strip() or not st.session_state.get("edu_deg_1","").strip():
        tips.append("Complete Education #1 with institute and degree.")
    return tips[:8]

with st.sidebar:
    st.header("💡 Suggestions")
    # Render up to 8 suggestions; otherwise show all-good message.
    tips = build_suggestions()
    if tips:
        for t in tips: st.write("• " + t)
    else:
        st.write("Looks good! ✅")

# ====================== SCORING (ATS / ROLE-FIT / QUALITY) ======================
# Use the model to produce a lightweight score and 2–4 reasons.
# Modes:
#  - ATS    : real JD provided
#  - SYNTH  : no JD, but a target role given → synthesize JD
#  - QUALITY: neither JD nor target role → generic resume quality

def ats_score(resume_text: str, jd: str, target_role: str = "", skills_text: str = ""):
    # Decide mode based on inputs
    mode = "ATS" if jd.strip() else ("SYNTH" if target_role.strip() else "QUALITY")

    # If synthesizing, build a tiny role spec for the model
    if mode == "SYNTH":
        jd = f"""
            Role: {target_role}
            Key skills to look for: {skills_text}
            Focus on: impact statements, internships/projects relevant to {target_role}, core tools listed above.
            """
    # Build strict-JSON prompts (short to save tokens)
    if mode in {"ATS","SYNTH"}:
        prompt = f"""
                You are an ATS assistant. Score the candidate resume against the job description from 0 to 100.
                Return STRICT JSON with keys: score (integer 0-100), reasons (array of brief strings, max 4).

                Resume:

                Job Description:

                Only return JSON.
                """
    else:
        prompt = f"""
                You are a resume quality checker. Score the resume from 0 to 100 on structure, completeness,
                action verbs, measurable outcomes, readability, and section coverage for entry-level tech roles.
                Return STRICT JSON with keys: score (integer 0-100), reasons (array of brief strings, max 4).

                Resume:

                Only return JSON.
                """
    # Call model; parse JSON; clamp score to [0, 100]
    try:
        r = client.chat_completion(
            messages=[{"role": "user", "content": prompt}],
            max_tokens=300, temperature=0.2,
        )
        txt = r.choices[0].message["content"]
        m = re.search(r"\{.*\}", txt, re.S)
        data = json.loads(m.group(0)) if m else json.loads(txt)
        score = int(data.get("score", 0)); score = max(0, min(100, score))
        reasons = [str(x) for x in data.get("reasons", [])][:4]
        return score, reasons, mode
    except Exception:
        return None, [], mode

# ----- Sidebar: show score if resume exists -----
with st.sidebar:
    st.header("📊 Score")
    # Only score after first generation to save API calls
    if st.session_state.get("final_text"):
        skills_text = " | ".join([
            st.session_state.get("skills_lang",""),
            st.session_state.get("skills_fw",""),
            st.session_state.get("skills_db",""),
            st.session_state.get("skills_cloud",""),
            st.session_state.get("skills_soft",""),
        ])
        with st.spinner("Scoring…"):
            s, reasons, mode = ats_score(
                st.session_state["final_text"],
                st.session_state.get("job_desc",""),
                target_role=st.session_state.get("target_role",""),
                skills_text=skills_text
            )
        if s is None:
            st.caption("Couldn’t compute score. Try again after editing.")
        else:
            label = "ATS Match" if mode == "ATS" else ("Role Fit (synthetic JD)" if mode == "SYNTH" else "Quality Score")
            st.progress(s / 100)
            st.write(f"**{label}: {s}/100**")
            for r in reasons: st.write("• " + r)
    else:
        st.caption("Generate a resume to see a score.")

# ====================== EXPORTS ======================
# Build PDF/DOCX/TXT bytes in-memory for Streamlit download buttons.

# Export resume to PDF (A4). Minimal styling; fonts depend on selected template.
def export_pdf(text, name):
    # 1) Clean text and split header vs body
    # 2) Choose fonts/sizes by template (Classic/Professional/Modern)
    # 3) Center header; draw thin dividers between sections
    # 4) Bold project/publication titles; write body with multi_cell

    from fpdf import FPDF
    import os

    safe = sanitize_plain_text(text)
    lines = safe.splitlines()

    # Parse and split header/body
    hdr = []
    body_start = 0
    for _ in range(4):
        if body_start < len(lines) and lines[body_start] and normalize_heading(lines[body_start]) not in HEADINGS:
            hdr.append(lines[body_start]); body_start += 1
        else:
            break
    body_lines = lines[body_start:]

    pdf = FPDF(unit="mm", format="A4")
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_text_color(0,0,0)

    # ---------- font + sizing per template ----------
    def has_calibri():
        return os.path.exists("calibri.ttf") and os.path.exists("calibrib.ttf")

    if template == "Classic":
        heading_size, body_size, line_gap = 12, 11, 6
        font_reg, font_bold = ("Times", "Times")
        divider_color, divider_width = (0, 0, 0), 0.2
    elif template == "Professional":
        heading_size, body_size, line_gap = 12, 11, 6
        font_reg, font_bold = ("Helvetica", "Helvetica")
        divider_color, divider_width = (160, 160, 160), 0.2
    else:  # Modern -> Calibri (or Arial fallback)
        heading_size, body_size, line_gap = 13, 11, 6
        if has_calibri():
            try:
                pdf.add_font("Calibri", "", "calibri.ttf", uni=True)
                pdf.add_font("Calibri", "B", "calibrib.ttf", uni=True)
                font_reg, font_bold = ("Calibri", "Calibri")
            except Exception:
                font_reg, font_bold = ("Arial", "Arial")
        else:
            font_reg, font_bold = ("Arial", "Arial")
        divider_color, divider_width = (180, 180, 180), 0.6

    def is_heading(s): return normalize_heading(s) in HEADINGS

    # Centered header (name/title/contacts)
    for i, h in enumerate(hdr):
        pdf.set_font(font_bold if i == 0 else font_reg, "B" if i == 0 else "", size=(heading_size+2 if i==0 else body_size))
        pdf.multi_cell(0, line_gap + (1 if i==0 else 0), h, align="C")
    pdf.ln(2)

    # Divider helper
    def draw_divider():
        pdf.set_draw_color(*divider_color)
        pdf.set_line_width(divider_width)
        y = pdf.get_y()
        pdf.line(10, y, 200, y)
        pdf.ln(3)
        pdf.set_draw_color(0,0,0)
        pdf.set_line_width(0.2)

    # Body writer (headings, bold titles, paragraphs)
    in_projects = False
    in_publications = False
    first_section_written = False
    for raw in body_lines:
        line = raw.strip()
        if not line:
            continue

        if is_heading(line):
            if first_section_written:
                draw_divider()
            first_section_written = True
            in_projects = (normalize_heading(line) == "PROJECTS")
            in_publications = (normalize_heading(line) == "PUBLICATIONS")
            section_heading_size = heading_size + (1 if template == "Modern" else 0)
            pdf.set_font(font_bold, "B", section_heading_size)
            pdf.cell(0, line_gap + 1, normalize_heading(line), ln=1)
            pdf.set_font(font_reg, "", body_size)
            continue

        if (in_projects or in_publications) and is_project_or_pub_title_line(line):
            pdf.set_font(font_bold, "B", body_size)
            pdf.multi_cell(0, line_gap, line)
            pdf.set_font(font_reg, "", body_size)
            continue

        pdf.multi_cell(0, line_gap, line)

    return pdf.output(dest="S").encode("latin-1","ignore")

# Export resume to DOCX (Calibri 11). Simple bold headings + plain paragraphs.
def export_docx(text, name):
    # 1) Clean text
    # 2) Write 1–4 centered header lines (name/title/contacts)
    # 3) For each section: bold heading, blank gap, then content
    # 4) Bold project/publication titles

    from docx import Document
    from docx.shared import Pt
    safe = sanitize_plain_text(text)
    doc = Document()
    style = doc.styles["Normal"]; style.font.name = "Calibri"; style.font.size = Pt(11)
    lines = safe.splitlines()
    
    def add_center(ptext, bold=False, size=14):
        p = doc.add_paragraph(); p.alignment = 1
        r = p.add_run(ptext); r.bold = bold; r.font.size = Pt(size)
    idx = 0
    for k in range(4):
        if idx < len(lines) and lines[idx] and normalize_heading(lines[idx]) not in HEADINGS:
            add_center(lines[idx], bold=(k==0), size=(16 if k==0 else 11)); idx += 1
        else: break
    # body with blank divider between sections (no decorative chars)
    in_projects = False
    in_publications = False
    first_section_written = False
    for raw in lines[idx:]:
        line = raw.strip()
        if not line:
            continue
        
        if normalize_heading(line) in HEADINGS:
            if first_section_written:
                doc.add_paragraph("")  # clean gap divider
            first_section_written = True
            in_projects = (normalize_heading(line)=="PROJECTS")
            in_publications = (normalize_heading(line)=="PUBLICATIONS")
            p = doc.add_paragraph(); r = p.add_run(normalize_heading(line)); r.bold=True; r.font.size=Pt(12)
            continue
        
        if (in_projects or in_publications) and is_project_or_pub_title_line(line):
            p = doc.add_paragraph(); r = p.add_run(line); r.bold = True; continue
        doc.add_paragraph(line)
    out = BytesIO(); doc.save(out); return out.getvalue()

# ====================== EDITOR + SAVE + DOWNLOADS ======================
# After generation, let user tweak text, save to session, and download files.

if st.session_state.get("final_text"):
    st.subheader("✏️ Edit before download")
    # Keep an editable copy; warn if there are unsaved changes.
    # Download buttons: PDF / DOCX / TXT
    
    st.session_state.setdefault("editor", st.session_state["final_text"])
    # Show editor and unsaved-change hint
    editor_val = st.text_area("Make quick fixes. Click **Save changes** to update.", key="editor", height=700)

    if editor_val != st.session_state["final_text"]:
        st.info("You have unsaved edits. Click **Save changes** to apply them.")

    # Persist edits back to session
    if st.button("💾 Save changes"):
        st.session_state["final_text"] = st.session_state["editor"]
        st.success("Changes saved. You can download now ✅")

    final_text = st.session_state["final_text"]
    c1, c2, c3 = st.columns(3)
    # Build filenames safely from name; sanitize text for TXT.
    # Use in-memory bytes for Streamlit's download buttons.
    c1.download_button(
        "⬇️ Download PDF",
        data=export_pdf(final_text, name),
        file_name=f"{(name or 'resume').replace(' ', '_')}.pdf",
        mime="application/pdf",
        use_container_width=True,
    )
    c2.download_button(
        "⬇️ Download DOCX",
        data=export_docx(final_text, name),
        file_name=f"{(name or 'resume').replace(' ', '_')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
    )
    c3.download_button(
        "⬇️ Download TXT",
        data=sanitize_plain_text(final_text).encode("utf-8"),
        file_name=f"{(name or 'resume').replace(' ', '_')}.txt",
        mime="text/plain",
        use_container_width=True,
    )

    # ====================== PORTFOLIO BUILDER ======================
    # Convert the resume sections into a simple 1-page portfolio site (index.html + styles.css).

    # Split final resume text into {SECTION: body} blocks and preserve order.
    def extract_section_blocks(full_text: str):
        blocks = {}
        order = []
        current = None
        lines = full_text.splitlines()
        # skip 1-4 header lines (name/location/contacts)
        i = 0; header_lines = 0
        while i < len(lines) and lines[i].strip() and normalize_heading(lines[i]) not in HEADINGS:
            header_lines += 1; i += 1
            if header_lines >= 4: break
        body = "\n".join(lines[i:])
        cur = None; buf = []
        for ln in body.splitlines():
            if normalize_heading(ln) in HEADINGS:
                if cur:
                    blocks[cur] = "\n".join(buf).strip()
                    buf = []
                cur = normalize_heading(ln); order.append(cur)
            else:
                buf.append(ln)
        if cur:
            blocks[cur] = "\n".join(buf).strip()
        return blocks, order

    # Convert bare URLs to clickable <a> links.
    def linkify(text: str) -> str:
        text = re.sub(r'\\1', '', text)  # remove escaped artifact if any
        return re.sub(r'(https?://[^\s]+)', r'<a href="\1" target="_blank">\1</a>', text)


    # Render HTML + CSS for the portfolio (Modern/Professional theme).
    def build_portfolio_html(full_text: str, name: str, pro_title: str, location: str,
                             email: str, phone: str, linkedin: str, github: str,
                             theme: str = "Modern"):
        blocks, order = extract_section_blocks(full_text)
        is_professional = (theme == "Professional")

        styles_css = f"""
:root {{
  --bg: #ffffff;
  --card: #ffffff;
  --text: {('#222222' if is_professional else '#111827')};
  --sub: {('#555555' if is_professional else '#4b5563')};
  --accent: {('#222222' if is_professional else '#111827')};
  --divider: {('#e6e6e6' if is_professional else '#e5e7eb')};
}}
html, body {{ background: var(--bg); color: var(--text); font-family: Inter, Arial, sans-serif; }}
/* Keep theme color but make links obviously clickable */
a {{ color: var(--accent); text-decoration: underline; font-weight: 600; }}
a:hover {{ text-decoration: underline; opacity: 0.9; }}
.sec-body a {{ text-decoration: underline; font-weight: 600; }}
.container {{ max-width: 900px; margin: 24px auto 80px; padding: 0 18px; }}
.hero {{ text-align: center; padding: 40px 16px 16px; border-bottom: 1px solid var(--divider); }}
.hero h1 {{ font-family: {("'Playfair Display', serif" if is_professional else "Inter, Arial, sans-serif")}; letter-spacing: .5px; font-size: 32px; font-weight: 800; color: var(--text); }}
.hero .title {{ margin-top: 6px; color: var(--sub); font-weight: 600; }}
.hero .loc {{ margin-top: 4px; color: var(--sub); }}
.hero .contacts {{ margin-top: 10px; color: var(--text); }}
.sec {{
  padding: 16px;
  margin: 14px 0;
  border-radius: {('0px' if is_professional else '16px')};
  background: {('transparent' if is_professional else 'var(--card)')};
  border: {('1px solid var(--divider)' if is_professional else '1px solid #eef2ff')};
  box-shadow: {('none' if is_professional else '0 6px 24px rgba(0,0,0,.08)')};
}}
.theme-professional .sec {{ border: none; border-bottom: 1px solid var(--divider); border-radius: 0; background: transparent; padding: 20px 0; }}
.theme-professional .sec:last-child {{ border-bottom: none; }}
.sec h2 {{ font-size: 18px; font-weight: 800; letter-spacing: .4px; margin-bottom: 8px; color: {('var(--text)' if is_professional else 'var(--accent)')}; }}
.sec-body {{ line-height: 1.6; color: var(--text); }}
.sec-body .li {{ margin-left: 12px; }}
.foot {{ text-align:center; color: var(--sub); padding: 24px 0 40px; border-top: 1px solid var(--divider); margin-top: 32px; }}
"""

        def sec_html(title_key):
            if title_key not in blocks or not blocks[title_key].strip():
                return ""
            lines = [ln.strip() for ln in blocks[title_key].splitlines() if ln.strip()]
            html_lines = []
            for ln in lines:
                if ln.startswith("- "):
                    html_lines.append(f"<div class='li'>• {linkify(ln[2:])}</div>")
                else:
                    # For POSITIONS header + next-line details style, keep plain divs
                    html_lines.append(f"<div>{linkify(ln)}</div>")
            body_html = "\n".join(html_lines)
            return f"""
            <section class="sec">
              <h2>{title_key.title()}</h2>
              <div class="sec-body">{body_html}</div>
            </section>
            """

        wanted = [
            "PROFESSIONAL OVERVIEW","EDUCATION","SKILLS","EXPERIENCE / INTERNSHIPS",
            "PROJECTS","PUBLICATIONS","CERTIFICATIONS / HANDS-ON","ACHIEVEMENTS","PARTICIPATIONS",
            "POSITIONS OF RESPONSIBILITY / CO-CURRICULAR INVOLVEMENT"
        ]
        html_sections = [sec_html(k) for k in wanted if k in blocks and blocks[k].strip()]

        index_html = f"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1" />
  <title>{(name or 'Portfolio')} – {pro_title or ''}</title>
  <style>*{{box-sizing:border-box;margin:0;padding:0}}</style>
  <link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
  <link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;600;700&family=Playfair+Display:wght@600;700&display=swap" rel="stylesheet">
  <style>{styles_css}</style>
  <link rel="stylesheet" href="styles.css" />
</head>
<body class="theme-{'professional' if is_professional else 'modern'}">
  <header class="hero">
    <h1>{(name or '').upper()}</h1>
    <div class="title">{pro_title or ''}</div>
    {f"<div class='loc'>{location}</div>" if location else ""}
    <div class="contacts">{ " | ".join([s for s in [f"<b>Email:</b> {email}" if email else "", f"<b>Phone:</b> {phone}" if phone else "", f"<b>LinkedIn:</b> <a href='"+linkedin+"' target='_blank'>"+linkedin+"</a>" if linkedin else "", f"<b>GitHub:</b> <a href='"+github+"' target='_blank'>"+github+"</a>" if github else ""] if s]) }</div>
  </header>

  <main class="container">
    {"".join(html_sections)}
  </main>

  <footer class="foot">
    © {datetime.now().year} {(name or '')}. Built with AI Resume & Portfolio Builder.
  </footer>
</body>
</html>
"""
        return index_html, styles_css

    # Package index.html + styles.css into a .zip (in-memory).
    def make_portfolio_zip(index_html: str, styles_css: str, filename: str = "portfolio_site.zip"):
        bio = BytesIO()
        with zipfile.ZipFile(bio, mode="w", compression=zipfile.ZIP_DEFLATED) as z:
            z.writestr("index.html", index_html)
            z.writestr("styles.css", styles_css)  # optional external file
        bio.seek(0)
        return bio.getvalue()

    # Quick gate: only enable Portfolio download when minimal resume data is present.
    def form_is_complete() -> bool:
        has_contact = all([name and name.strip(), email and email.strip(), phone and phone.strip(), location and location.strip()])
        has_core = any([skills_lang.strip(), skills_fw.strip(), skills_db.strip(), skills_cloud.strip()])
        has_edu = bool((st.session_state.get("edu_inst_1","").strip() and st.session_state.get("edu_deg_1","").strip()))
        has_proj_or_exp = any([
            st.session_state.get("proj_title_1","").strip(),
            st.session_state.get("exp_comp_1","").strip() and st.session_state.get("exp_role_1","").strip(),
        ])
        return has_contact and has_core and has_edu and has_proj_or_exp

    # If form complete, generate portfolio site and expose as .zip download.
    # Otherwise, show guidance on what fields to fill.
    if st.session_state.get("final_text"):
        final_text = st.session_state["final_text"]
        if form_is_complete():
            idx_html, css_txt = build_portfolio_html(
                final_text,
                name=name, pro_title=pro_title, location=location,
                email=email, phone=phone, linkedin=linkedin, github=github,
                theme=portfolio_theme
            )
            zip_bytes = make_portfolio_zip(idx_html, css_txt)
            st.download_button(
                "🌐 Download Portfolio (HTML .zip)",
                data=zip_bytes,
                file_name=f"{(name or 'portfolio').replace(' ', '_')}_site.zip",
                mime="application/zip",
                use_container_width=True,
            )
        else:
            st.info("Fill all required fields (contact, basic skills, education #1, and at least one project or experience) and generate your resume to enable the Portfolio download.")

# ====================== END OF FILE ======================
# Notes for contributors:
# - Keep comments short, action-style (one line above logic blocks)
# - Follow section dividers exactly (helps readability in large files)
# - Avoid printing secrets; rely on environment variables (HF_TOKEN, etc.)
# - Before committing: run Streamlit once locally to ensure no syntax breaks
# - Feel free to swap in any chat-instruct model via HF_MODEL env var